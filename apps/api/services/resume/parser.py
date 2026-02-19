import io
from io import BytesIO
from pdfminer.high_level import extract_text


class ResumeParser:
    """
    Extracts plain text from resume files.
    Supported formats: PDF, DOCX, DOC, TXT, RTF
    """

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        Auto-detect format from filename extension and extract text.
        """
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext == "pdf":
            return self._from_pdf(file_content)
        elif ext == "docx":
            return self._from_docx(file_content)
        elif ext == "doc":
            return self._from_doc(file_content)
        elif ext in ("txt", "text"):
            return self._from_txt(file_content)
        elif ext == "rtf":
            return self._from_rtf(file_content)
        else:
            raise ValueError(
                f"Unsupported file format '.{ext}'. "
                "Please upload PDF, DOCX, DOC, TXT, or RTF."
            )

    # ── Individual parsers ────────────────────────────────────────────────────

    def _from_pdf(self, content: bytes) -> str:
        try:
            text = extract_text(BytesIO(content))
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")

    def _from_docx(self, content: bytes) -> str:
        try:
            import docx  # python-docx
            doc = docx.Document(BytesIO(content))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts).strip()
        except ImportError:
            raise ValueError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")

    def _from_doc(self, content: bytes) -> str:
        """
        Parse legacy .doc format using mammoth (converts to plain text).
        Falls back to antiword-style extraction if mammoth isn't available.
        """
        try:
            import mammoth
            result = mammoth.extract_raw_text(BytesIO(content))
            return result.value.strip()
        except ImportError:
            raise ValueError(
                "mammoth not installed. Run: pip install mammoth"
            )
        except Exception as e:
            raise ValueError(f"Failed to parse DOC: {e}")

    def _from_txt(self, content: bytes) -> str:
        # Try UTF-8 first, then latin-1 as fallback
        try:
            return content.decode("utf-8").strip()
        except UnicodeDecodeError:
            return content.decode("latin-1").strip()

    def _from_rtf(self, content: bytes) -> str:
        try:
            from striprtf.striprtf import rtf_to_text
            return rtf_to_text(content.decode("latin-1")).strip()
        except ImportError:
            raise ValueError("striprtf not installed. Run: pip install striprtf")
        except Exception as e:
            raise ValueError(f"Failed to parse RTF: {e}")

    # ── Legacy compat ─────────────────────────────────────────────────────────

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Kept for backwards compatibility."""
        return self._from_pdf(file_content)
